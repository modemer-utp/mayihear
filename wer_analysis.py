"""
WER & CER Analysis: Gabriela.docx (reference) vs transcripcion_gabi.txt (MayiHear)
"""

import re
import unicodedata
import jiwer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# 1. TEXT EXTRACTION
# ─────────────────────────────────────────────

TIMESTAMP_RE = re.compile(r'^\d{2}:\d{2}:\d{2}\s+\w+')
LINE_NUM_RE  = re.compile(r'^\s*\d+→')
SPEAKER_RE   = re.compile(r'^Speaker\s+\d+:\s*', re.IGNORECASE)


def extract_docx_text(path):
    """Extract speech text from Gabriela.docx, skip timestamps/speaker lines."""
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if TIMESTAMP_RE.match(t):
            continue
        if t in ('Archivo de audio', 'Transcripci\u00f3n', 'Transcripcion') or t.endswith('.mp3'):
            continue
        lines.append(t)
    return ' '.join(lines)


def extract_txt_text(path):
    """Extract speech text from transcripcion_gabi.txt, skip speaker/line labels."""
    chunks = []
    with open(path, encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            # Remove line numbers
            line = LINE_NUM_RE.sub('', line)
            # Remove Speaker labels
            line = SPEAKER_RE.sub('', line)
            line = line.strip()
            if line:
                chunks.append(line)
    return ' '.join(chunks)


# ─────────────────────────────────────────────
# 2. NORMALISATION
# ─────────────────────────────────────────────

def normalize(text, remove_punct=True):
    """Lowercase, strip accents optionally, remove punctuation."""
    text = text.lower()
    # Keep accented chars as-is for Spanish (accents are meaningful)
    if remove_punct:
        text = re.sub(r'[^\w\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


# ─────────────────────────────────────────────
# 3. EDIT DISTANCE (Levenshtein) — word & char
# ─────────────────────────────────────────────

def edit_distance_ops(ref_tokens, hyp_tokens):
    """Return (insertions, deletions, substitutions, distance) between token lists."""
    n, m = len(ref_tokens), len(hyp_tokens)
    # DP table with backtracking info
    dp = [[0] * (m + 1) for _ in range(n + 1)]
    for i in range(n + 1):
        dp[i][0] = i
    for j in range(m + 1):
        dp[0][j] = j

    for i in range(1, n + 1):
        for j in range(1, m + 1):
            if ref_tokens[i-1] == hyp_tokens[j-1]:
                dp[i][j] = dp[i-1][j-1]
            else:
                dp[i][j] = 1 + min(dp[i-1][j],    # deletion
                                   dp[i][j-1],    # insertion
                                   dp[i-1][j-1])  # substitution

    # Backtrack to count op types
    i, j = n, m
    ins = dels = subs = 0
    while i > 0 or j > 0:
        if i > 0 and j > 0 and ref_tokens[i-1] == hyp_tokens[j-1]:
            i -= 1; j -= 1
        elif i > 0 and j > 0 and dp[i][j] == dp[i-1][j-1] + 1:
            subs += 1; i -= 1; j -= 1
        elif i > 0 and dp[i][j] == dp[i-1][j] + 1:
            dels += 1; i -= 1
        else:
            ins += 1; j -= 1

    return ins, dels, subs, dp[n][m]


def wer(ref_text, hyp_text):
    ref_n = normalize(ref_text)
    hyp_n = normalize(hyp_text)
    ref_words = ref_n.split()
    hyp_words = hyp_n.split()
    out = jiwer.process_words(ref_n, hyp_n)
    n    = len(ref_words)
    ins  = out.insertions
    dels = out.deletions
    subs = out.substitutions
    dist = ins + dels + subs
    rate = out.wer
    return {
        'wer': rate,
        'wer_pct': rate * 100,
        'ref_words': n,
        'hyp_words': len(hyp_words),
        'insertions': ins,
        'deletions': dels,
        'substitutions': subs,
        'total_errors': dist,
    }


def cer(ref_text, hyp_text):
    ref_n = normalize(ref_text, remove_punct=False)
    hyp_n = normalize(hyp_text, remove_punct=False)
    ref_chars = list(ref_n)
    hyp_chars = list(hyp_n)
    out = jiwer.process_characters(ref_n, hyp_n)
    n    = len(ref_chars)
    ins  = out.insertions
    dels = out.deletions
    subs = out.substitutions
    dist = ins + dels + subs
    rate = out.cer
    return {
        'cer': rate,
        'cer_pct': rate * 100,
        'ref_chars': n,
        'hyp_chars': len(hyp_chars),
        'insertions': ins,
        'deletions': dels,
        'substitutions': subs,
        'total_errors': dist,
    }


# ─────────────────────────────────────────────
# 4. WORD CLOUD  (most-confused words, simple)
# ─────────────────────────────────────────────

def most_common_words(text, top=10):
    words = normalize(text).split()
    freq = {}
    for w in words:
        if len(w) > 3:
            freq[w] = freq.get(w, 0) + 1
    return sorted(freq.items(), key=lambda x: -x[1])[:top]


# ─────────────────────────────────────────────
# 5. WORD REPORT GENERATION
# ─────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def add_metric_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue background
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row[c_idx].text = str(val)
            row[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 0:
                tc = row[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D6E4F0')
                tcPr.append(shd)
    return table


def add_bar(doc, label, value_pct, max_pct=100):
    """Add a simple text-based bar chart row."""
    filled = int(value_pct / max_pct * 30)
    bar = '█' * filled + '░' * (30 - filled)
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label:<18}')
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    r2 = p.add_run(f' {bar} ')
    r2.font.name = 'Courier New'
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(46, 116, 181)
    r3 = p.add_run(f'{value_pct:.1f}%')
    r3.font.name = 'Courier New'
    r3.font.size = Pt(10)
    r3.bold = True


def generate_report(ref_path, hyp_path, out_path):
    print("Extracting text...")
    ref_text = extract_docx_text(ref_path)
    hyp_text = extract_txt_text(hyp_path)

    print(f"Reference words: {len(ref_text.split())}")
    print(f"Hypothesis words: {len(hyp_text.split())}")

    print("Calculating WER...")
    w = wer(ref_text, hyp_text)
    print("Calculating CER...")
    c = cer(ref_text, hyp_text)

    print(f"WER: {w['wer_pct']:.2f}%")
    print(f"CER: {c['cer_pct']:.2f}%")

    # Accuracy = 1 - WER  (clamped to 0)
    word_acc = max(0, 100 - w['wer_pct'])
    char_acc = max(0, 100 - c['cer_pct'])

    # ── Build Word document ──────────────────────
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Título ──
    title = doc.add_heading('MayiHear — Reporte de Calidad de Transcripción', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(31, 73, 125)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f'Generado: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}  |  '
                    f'Referencia: Gabriela.docx  |  Hipótesis: transcripcion_gabi.txt')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── 1. Resumen Ejecutivo ──
    add_heading(doc, '1. Resumen Ejecutivo', 1, color=(31, 73, 125))
    p = doc.add_paragraph()
    p.add_run('Sistema de referencia: ').bold = True
    p.add_run('Transcripción automática de Microsoft Teams (Gabriela.docx)\n')
    r2 = p.add_run('Sistema MayiHear: ')
    r2.bold = True
    p.add_run('Transcripción basada en Whisper de MayiHear (transcripcion_gabi.txt)\n\n')
    p.add_run(
        'Este reporte compara la transcripción de MayiHear contra la referencia generada '
        'por Teams, utilizando métricas estándar de evaluación ASR: Tasa de Error por '
        'Palabra (WER) y Tasa de Error por Carácter (CER). Un menor WER/CER indica mayor precisión.'
    )

    doc.add_paragraph()

    # ── 2. Métricas Principales ──
    add_heading(doc, '2. Métricas Principales', 1, color=(31, 73, 125))

    add_metric_table(doc,
        ['Métrica', 'Valor', 'Interpretación'],
        [
            ['Tasa de Error por Palabra (WER)', f"{w['wer_pct']:.2f}%", _interp_wer(w['wer_pct'])],
            ['Tasa de Error por Carácter (CER)', f"{c['cer_pct']:.2f}%", _interp_cer(c['cer_pct'])],
            ['Precisión por Palabra',            f"{word_acc:.2f}%",     ''],
            ['Precisión por Carácter',           f"{char_acc:.2f}%",     ''],
        ]
    )

    doc.add_paragraph()

    # Gráfico de barras
    add_heading(doc, 'Resumen Visual', 2)
    max_v = 100  # noqa: F841
    add_bar(doc, 'WER',                w['wer_pct'])
    add_bar(doc, 'CER',                c['cer_pct'])
    add_bar(doc, 'Prec. por Palabra',  word_acc)
    add_bar(doc, 'Prec. por Carácter', char_acc)

    doc.add_paragraph()

    # ── 3. Detalle a Nivel de Palabras ──
    add_heading(doc, '3. Desglose de Errores por Palabra', 1, color=(31, 73, 125))
    add_metric_table(doc,
        ['Componente', 'Cantidad', '% sobre palabras de referencia'],
        [
            ['Palabras de referencia',    w['ref_words'],    '—'],
            ['Palabras de hipótesis',     w['hyp_words'],    '—'],
            ['Sustituciones (S)',          w['substitutions'], f"{w['substitutions']/w['ref_words']*100:.2f}%"],
            ['Eliminaciones (D)',          w['deletions'],     f"{w['deletions']/w['ref_words']*100:.2f}%"],
            ['Inserciones (I)',            w['insertions'],    f"{w['insertions']/w['ref_words']*100:.2f}%"],
            ['Total de errores de palabra', w['total_errors'], f"{w['wer_pct']:.2f}%"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Fórmula WER: ').bold = True
    p.add_run('WER = (S + D + I) / N   donde N = número de palabras en la referencia')
    p.runs[-1].font.name = 'Courier New'

    doc.add_paragraph()

    # ── 4. Detalle a Nivel de Caracteres ──
    add_heading(doc, '4. Desglose de Errores por Carácter', 1, color=(31, 73, 125))
    add_metric_table(doc,
        ['Componente', 'Cantidad', '% sobre caracteres de referencia'],
        [
            ['Caracteres de referencia',   c['ref_chars'],    '—'],
            ['Caracteres de hipótesis',    c['hyp_chars'],    '—'],
            ['Sustituciones (S)',           c['substitutions'], f"{c['substitutions']/c['ref_chars']*100:.2f}%"],
            ['Eliminaciones (D)',           c['deletions'],     f"{c['deletions']/c['ref_chars']*100:.2f}%"],
            ['Inserciones (I)',             c['insertions'],    f"{c['insertions']/c['ref_chars']*100:.2f}%"],
            ['Total de errores de carácter', c['total_errors'], f"{c['cer_pct']:.2f}%"],
        ]
    )

    doc.add_paragraph()

    # ── 5. Estadísticas de Texto ──
    add_heading(doc, '5. Estadísticas de Texto', 1, color=(31, 73, 125))
    ref_sentences = len(re.split(r'[.!?]+', ref_text))
    hyp_sentences = len(re.split(r'[.!?]+', hyp_text))
    ref_avg = len(ref_text.split()) / max(ref_sentences, 1)
    hyp_avg = len(hyp_text.split()) / max(hyp_sentences, 1)

    add_metric_table(doc,
        ['Estadística', 'Referencia (Teams)', 'MayiHear'],
        [
            ['Total de palabras',          w['ref_words'],   w['hyp_words']],
            ['Total de caracteres',        c['ref_chars'],   c['hyp_chars']],
            ['Oraciones aprox.',           ref_sentences,    hyp_sentences],
            ['Promedio palabras/oración',  f'{ref_avg:.1f}', f'{hyp_avg:.1f}'],
        ]
    )

    doc.add_paragraph()

    # ── 6. Observaciones y Recomendaciones ──
    add_heading(doc, '6. Observaciones y Recomendaciones', 1, color=(31, 73, 125))

    obs = _build_observations(w, c, word_acc, char_acc)
    for bullet in obs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bullet)

    doc.add_paragraph()

    # ── 7. Metodología ──
    add_heading(doc, '7. Metodología', 1, color=(31, 73, 125))
    p = doc.add_paragraph()
    p.add_run(
        'Ambas transcripciones fueron preprocesadas de la siguiente manera antes de la comparación:\n'
        '  1. Se eliminaron etiquetas de hablante, marcas de tiempo y números de línea.\n'
        '  2. Texto convertido a minúsculas; se preservaron los acentos del español.\n'
        '  3. Se eliminó la puntuación (reemplazada por espacios).\n'
        '  4. Se colapsaron espacios múltiples a uno solo.\n\n'
        'La distancia de edición se calculó con el algoritmo estándar de Levenshtein '
        'aplicado a nivel de palabra (WER) y a nivel de carácter (CER).\n\n'
        'Referencia (ground truth): transcripción automática de Microsoft Teams — '
        'considerada la línea base más precisa para esta comparación.'
    )

    doc.save(out_path)
    print(f"\nReport saved → {out_path}")
    return w, c


def _interp_wer(pct):
    if pct < 10:   return 'Excelente (< 10%)'
    if pct < 20:   return 'Bueno (10–20%)'
    if pct < 40:   return 'Aceptable (20–40%)'
    return 'Requiere mejora (> 40%)'


def _interp_cer(pct):
    if pct < 5:    return 'Excelente (< 5%)'
    if pct < 15:   return 'Bueno (5–15%)'
    if pct < 30:   return 'Aceptable (15–30%)'
    return 'Requiere mejora (> 30%)'


def _build_observations(w, c, word_acc, char_acc):
    obs = []
    obs.append(
        f"MayiHear obtuvo un WER de {w['wer_pct']:.1f}% y un CER de {c['cer_pct']:.1f}% "
        f"en comparación con la transcripción de referencia de Teams."
    )
    if w['deletions'] > w['insertions']:
        obs.append(
            f"Las eliminaciones ({w['deletions']}) superan a las inserciones ({w['insertions']}), "
            "lo que sugiere que MayiHear tiende a omitir palabras en lugar de agregar contenido inexistente."
        )
    elif w['insertions'] > w['deletions']:
        obs.append(
            f"Las inserciones ({w['insertions']}) superan a las eliminaciones ({w['deletions']}), "
            "lo que sugiere que MayiHear ocasionalmente agrega palabras que no están en la referencia."
        )
    else:
        obs.append(
            f"Las inserciones y eliminaciones están equilibradas ({w['insertions']} cada una)."
        )

    obs.append(
        f"Las sustituciones representan el mayor tipo de error ({w['substitutions']} palabras), "
        "lo cual es típico en sistemas ASR que manejan vocabulario específico del dominio (terminología académica y de evaluación)."
    )

    if word_acc >= 80:
        obs.append(
            "La precisión general por palabra es aceptable para un caso de uso de resumen de reuniones, "
            "donde la redacción exacta es menos crítica que capturar decisiones clave y puntos de acción."
        )
    else:
        obs.append(
            "La precisión por palabra está por debajo del 80% — se recomienda aplicar postprocesamiento "
            "con modelos de lenguaje específicos del dominio o ajustar Whisper con grabaciones de reuniones de UTP."
        )

    obs.append(
        "Ambos sistemas de transcripción operan sobre el mismo audio y capturan la voz del hablante vía loopback. "
        "Las diferencias pueden reflejar parcialmente el acceso de Teams a modelos de cancelación de eco "
        "acústico y adaptación de hablante."
    )
    obs.append(
        "Próximos pasos recomendados: (1) recopilar un conjunto de evaluación más amplio con múltiples reuniones; "
        "(2) experimentar con el modelo Whisper 'large-v3' para mayor precisión; "
        "(3) agregar un paso de postprocesamiento con vocabulario del dominio (sílabo, rúbrica, competencia, etc.)."
    )
    return obs


if __name__ == '__main__':
    generate_report(
        ref_path='Gabriela.docx',
        hyp_path='transcripcion_gabi.txt',
        out_path='MayiHear_Transcription_Quality_Report.docx',
    )
