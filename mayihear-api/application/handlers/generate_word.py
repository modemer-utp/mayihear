from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from domain.models.output.meeting_act_result import MeetingActResult


def build_word_document(act: MeetingActResult) -> BytesIO:
    doc = Document()

    # ── Title ────────────────────────────────────────────────────
    title = doc.add_heading(act.nombre_reunion or "Acta de Reunión", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 1. Datos Básicos ─────────────────────────────────────────
    doc.add_heading("1. Datos Básicos", level=2)

    doc.add_paragraph(f"Fecha: {act.fecha}")

    if act.participantes:
        p = doc.add_paragraph("Participantes: ")
        p.runs[0].bold = True
        p.add_run(", ".join(act.participantes))

    # ── 2. Contenido de la Reunión ───────────────────────────────
    doc.add_heading("2. Contenido de la Reunión", level=2)

    if act.resumen_ejecutivo:
        doc.add_heading("Resumen Ejecutivo", level=3)
        doc.add_paragraph(act.resumen_ejecutivo)

    if act.temas:
        doc.add_heading("Temas Tratados", level=3)
        for topic in act.temas:
            doc.add_heading(topic.titulo, level=4)
            if topic.avances:
                doc.add_paragraph("Avances:", style="List Bullet").runs[0].bold = True
                for a in topic.avances:
                    doc.add_paragraph(a, style="List Bullet 2")
            if topic.bloqueantes:
                doc.add_paragraph("Bloqueantes:", style="List Bullet").runs[0].bold = True
                for b in topic.bloqueantes:
                    doc.add_paragraph(b, style="List Bullet 2")
            if topic.aprendizajes:
                doc.add_paragraph("Aprendizajes:", style="List Bullet").runs[0].bold = True
                for ap in topic.aprendizajes:
                    doc.add_paragraph(ap, style="List Bullet 2")

    if act.acuerdos:
        doc.add_heading("Acuerdos y Compromisos", level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Acción"
        hdr[1].text = "Responsable"
        for h in hdr:
            for run in h.paragraphs[0].runs:
                run.bold = True
        for acuerdo in act.acuerdos:
            row = table.add_row().cells
            row[0].text = acuerdo.accion
            row[1].text = acuerdo.responsable or "—"

    # ── 3. Aspectos Adicionales ──────────────────────────────────
    doc.add_heading("3. Aspectos Adicionales", level=2)

    if act.riesgos:
        doc.add_heading("Riesgos Identificados", level=3)
        for r in act.riesgos:
            doc.add_paragraph(r, style="List Bullet")

    if act.pendientes_reunion_anterior:
        doc.add_heading("Pendientes de Reunión Anterior", level=3)
        for p in act.pendientes_reunion_anterior:
            doc.add_paragraph(p, style="List Bullet")

    if act.proxima_reunion:
        doc.add_heading("Próxima Reunión", level=3)
        doc.add_paragraph(act.proxima_reunion)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
