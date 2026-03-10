"""
Genera un reporte de costos de MayiHear en Word (.docx)
Ejecutar desde: D:\Proyectos actuales\mayihear-utp\
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from datetime import datetime

LOG_PATH = r"D:\Proyectos actuales\mayihear-utp\mayihear-api\data\usage_log.json"
OUTPUT   = r"D:\Proyectos actuales\mayihear-utp\reporte_costos_mayihear.docx"

# ── Colores ──────────────────────────────────────────────────
DARK     = RGBColor(0x1e, 0x29, 0x3b)
BLUE     = RGBColor(0x1d, 0x4e, 0xd8)
GRAY     = RGBColor(0x64, 0x74, 0x8b)
GREEN    = RGBColor(0x05, 0x96, 0x69)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1e293b"
ROW_ALT   = "f8fafc"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:color'), kwargs.get('color', 'e2e8f0'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def bold_run(para, text, size=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def normal_run(para, text, size=11, color=None):
    run = para.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def fmt_dur(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    if h:
        return f"{h}h {m:02d}min"
    return f"{m}min"

def fmt_usd(v):
    return f"${v:.4f}"

# ── Carga datos ───────────────────────────────────────────────
with open(LOG_PATH, encoding="utf-8") as f:
    records = json.load(f)

# Agrupar por sesión (mismo día, cerca en tiempo)
sessions = []
current  = []
for r in records:
    if not current:
        current.append(r)
    else:
        prev_ts  = datetime.fromisoformat(current[-1]["timestamp"])
        this_ts  = datetime.fromisoformat(r["timestamp"])
        diff_min = (this_ts - prev_ts).total_seconds() / 60
        if diff_min < 60:
            current.append(r)
        else:
            sessions.append(current)
            current = [r]
if current:
    sessions.append(current)

# ── Crea documento ────────────────────────────────────────────
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Portada / encabezado ──────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
bold_run(title_para, "MayiHear", 26, DARK)
normal_run(title_para, "  |  Reporte de Costos de IA", 16, GRAY)

sub_para = doc.add_paragraph()
normal_run(sub_para, f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}", 10, GRAY)

doc.add_paragraph()

# ── Resumen ejecutivo ─────────────────────────────────────────
total_cost  = sum(r["estimated_cost_usd"] for r in records)
trans_cost  = sum(r["estimated_cost_usd"] for r in records if r["type"] == "transcription")
ins_cost    = sum(r["estimated_cost_usd"] for r in records if r["type"] == "insights")
act_cost    = sum(r["estimated_cost_usd"] for r in records if r["type"] == "meeting_act")
total_dur   = sum(r.get("recording_duration_seconds", 0) for r in records if r["type"] == "transcription")

exec_h = doc.add_heading("Resumen Ejecutivo", level=1)
exec_h.runs[0].font.color.rgb = DARK

summary_table = doc.add_table(rows=2, cols=4)
summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
summary_table.style = "Table Grid"

headers = ["Reuniones procesadas", "Duración total grabada", "Costo total (USD)", "Costo promedio/reunión"]
values  = [
    str(len(sessions)),
    fmt_dur(total_dur),
    fmt_usd(total_cost),
    fmt_usd(total_cost / len(sessions))
]

for i, (h, v) in enumerate(zip(headers, values)):
    hcell = summary_table.rows[0].cells[i]
    vcell = summary_table.rows[1].cells[i]
    set_cell_bg(hcell, HEADER_BG)
    set_cell_bg(vcell, "f0f9ff")
    hp = hcell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(h)
    hr.font.color.rgb = WHITE
    hr.font.size = Pt(9)
    hr.bold = True
    vp = vcell.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(v)
    vr.font.size = Pt(14)
    vr.bold = True
    vr.font.color.rgb = BLUE

doc.add_paragraph()

# ── Desglose por sesión ───────────────────────────────────────
sess_h = doc.add_heading("Desglose por Sesión", level=1)
sess_h.runs[0].font.color.rgb = DARK

TYPE_LABELS = {
    "transcription": "Transcripción",
    "insights":      "Insights",
    "meeting_act":   "Acta de reunión",
}

for idx, session in enumerate(sessions):
    ts       = datetime.fromisoformat(session[0]["timestamp"])
    sess_dur = sum(r.get("recording_duration_seconds", 0) for r in session if r["type"] == "transcription")
    sess_tot = sum(r["estimated_cost_usd"] for r in session)

    # Subtítulo sesión
    sp = doc.add_paragraph()
    bold_run(sp, f"Sesión {idx + 1}", 12, DARK)
    normal_run(sp, f"   {ts.strftime('%d/%m/%Y')}   ·   {fmt_dur(sess_dur)}   ·   Total: {fmt_usd(sess_tot)}", 11, GRAY)

    # Tabla de sesión
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Inches(1.8), Inches(1.6), Inches(1.1), Inches(1.1), Inches(1.1)]
    col_headers = ["Paso", "Modelo", "Tokens entrada", "Tokens salida", "Costo (USD)"]

    hrow = tbl.rows[0]
    for i, (ch, cw) in enumerate(zip(col_headers, col_widths)):
        cell = hrow.cells[i]
        cell.width = cw
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        r = p.add_run(ch)
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)
        r.bold = True

    for ri, rec in enumerate(session):
        row = tbl.add_row()
        bg  = ROW_ALT if ri % 2 == 0 else "ffffff"
        cells_data = [
            TYPE_LABELS.get(rec["type"], rec["type"]),
            rec["model"],
            f"{rec['input_tokens']:,}",
            f"{rec['output_tokens']:,}",
            fmt_usd(rec["estimated_cost_usd"]),
        ]
        for ci, (val, cw) in enumerate(zip(cells_data, col_widths)):
            cell = row.cells[ci]
            cell.width = cw
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(10)
            if ci == 4:
                r.bold = True
                r.font.color.rgb = GREEN

    # Fila total sesión
    total_row = tbl.add_row()
    set_cell_bg(total_row.cells[0], "e2e8f0")
    set_cell_bg(total_row.cells[1], "e2e8f0")
    set_cell_bg(total_row.cells[2], "e2e8f0")
    set_cell_bg(total_row.cells[3], "e2e8f0")
    set_cell_bg(total_row.cells[4], "dcfce7")
    for ci in range(4):
        p = total_row.cells[ci].paragraphs[0]
        r = p.add_run("TOTAL SESIÓN" if ci == 0 else "")
        r.font.size = Pt(10); r.bold = True
    tp = total_row.cells[4].paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(fmt_usd(sess_tot))
    tr.font.size = Pt(11); tr.bold = True
    tr.font.color.rgb = GREEN

    doc.add_paragraph()

# ── Desglose por tipo ─────────────────────────────────────────
type_h = doc.add_heading("Desglose por Tipo de Operación", level=1)
type_h.runs[0].font.color.rgb = DARK

type_tbl = doc.add_table(rows=1, cols=3)
type_tbl.style = "Table Grid"
for i, h in enumerate(["Operación", "Costo total (USD)", "% del total"]):
    cell = type_tbl.rows[0].cells[i]
    set_cell_bg(cell, HEADER_BG)
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.color.rgb = WHITE; r.font.size = Pt(9); r.bold = True

for label, cost in [("Transcripción (gemini-2.5-pro)", trans_cost),
                    ("Insights (gemini-2.5-flash)", ins_cost),
                    ("Acta de reunión (gemini-2.5-flash)", act_cost)]:
    row = type_tbl.add_row()
    for ci, val in enumerate([label, fmt_usd(cost), f"{cost/total_cost*100:.1f}%"]):
        p = row.cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(10)
        if ci > 0: r.bold = True; r.font.color.rgb = GREEN

total_row = type_tbl.add_row()
set_cell_bg(total_row.cells[0], "e2e8f0")
set_cell_bg(total_row.cells[1], "dcfce7")
set_cell_bg(total_row.cells[2], "e2e8f0")
for ci, val in enumerate(["TOTAL", fmt_usd(total_cost), "100%"]):
    p = total_row.cells[ci].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(val)
    r.font.size = Pt(11); r.bold = True
    if ci == 1: r.font.color.rgb = GREEN

doc.add_paragraph()

# ── Estimación por duración ───────────────────────────────────
est_h = doc.add_heading("Estimación de Costos por Duración de Reunión", level=1)
est_h.runs[0].font.color.rgb = DARK

note = doc.add_paragraph()
normal_run(note, "Basado en el costo promedio de las sesiones registradas (gemini-2.5-pro para transcripción + gemini-2.5-flash para insights y acta).", 10, GRAY)

cost_per_sec = total_cost / total_dur if total_dur else 0

est_tbl = doc.add_table(rows=1, cols=3)
est_tbl.style = "Table Grid"
for i, h in enumerate(["Duración de reunión", "Costo estimado (USD)", "Costo estimado (S/.)"]):
    cell = est_tbl.rows[0].cells[i]
    set_cell_bg(cell, HEADER_BG)
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.color.rgb = WHITE; r.font.size = Pt(9); r.bold = True

PEN_RATE = 3.75
for minutes in [30, 45, 60, 90, 120, 150, 180]:
    est = cost_per_sec * minutes * 60
    row = est_tbl.add_row()
    for ci, val in enumerate([f"{minutes} minutos", fmt_usd(est), f"S/. {est * PEN_RATE:.2f}"]):
        p = row.cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(10)
        if ci > 0: r.bold = True; r.font.color.rgb = GREEN

doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
normal_run(footer, "MayiHear — UTP MVP  ·  Reporte generado automáticamente desde usage_log.json", 9, GRAY)

doc.save(OUTPUT)
print(f"Reporte generado: {OUTPUT}")
